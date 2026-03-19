"""
Utility functions for file I/O, JSON handling, timestamps, and logging.
"""

import json
import os
import re
import sys
from datetime import datetime
from typing import Any, Optional, List

import config


def log_info(message: str) -> None:
    """
    Print an informational log message with timestamp.

    Args:
        message (str): Message to log.
    """
    ts = datetime.now().strftime(config.LOG_DATETIME_FORMAT)
    print(f"[INFO] [{ts}] {message}")


def log_warning(message: str) -> None:
    """
    Print a warning log message with timestamp.

    Args:
        message (str): Message to log.
    """
    ts = datetime.now().strftime(config.LOG_DATETIME_FORMAT)
    print(f"[WARN] [{ts}] {message}", file=sys.stderr)


def log_error(message: str) -> None:
    """
    Print an error log message with timestamp.

    Args:
        message (str): Message to log.
    """
    ts = datetime.now().strftime(config.LOG_DATETIME_FORMAT)
    print(f"[ERROR] [{ts}] {message}", file=sys.stderr)


def current_timestamp() -> str:
    """
    Get the current timestamp as a compact string (useful for filenames).

    Returns:
        str: Timestamp string.
    """
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def read_text_file(path: str) -> str:
    """
    Read a UTF-8 text file and return its contents.

    Uses errors='replace' so that profiles with occasional non-UTF-8 bytes
    (e.g. curly quotes from Word exports) don't crash the pipeline.

    Args:
        path (str): Path to the text file.

    Returns:
        str: File contents.

    Raises:
        OSError: If the file cannot be read.
    """
    # FIX: errors="replace" prevents hard crash on malformed bytes while still
    # reading the vast majority of the file content correctly.
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def read_document_text(path: str, label: str = "input") -> str:
    """
    Read unstructured text from supported document types.

    Supported:
    - .txt  : UTF-8 text
    - .docx : extracted via python-docx
    - .pdf  : extracted via pdfplumber

    This function enables the pipeline to accept profile/JD inputs that are
    provided as documents, not just plain text.

    Args:
        path: File path to read.
        label: Human-friendly label for error messages (e.g., "profile", "jd").

    Returns:
        Extracted text as a single string.

    Raises:
        RuntimeError: If required extraction dependencies are missing.
        ValueError: If the file extension is unsupported.
        OSError: If the file cannot be read.
    """
    _, ext = os.path.splitext(path)
    ext = ext.lower()

    if ext in (".txt", ".text", ".md"):
        return read_text_file(path)
    if ext == ".docx":
        return _extract_docx_text(path, label=label)
    if ext == ".pdf":
        return _extract_pdf_text(path, label=label)

    raise ValueError(
        f"Unsupported {label} file type '{ext}'. Supported extensions: .txt, .docx, .pdf"
    )


def _extract_docx_text(docx_path: str, label: str = "input") -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        docx_path: Path to the DOCX file.
        label: Label for error messages.

    Returns:
        Extracted text.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "DOCX input support requires 'python-docx'. Install with: pip install -r requirements.txt"
        ) from exc

    doc = Document(docx_path)
    parts: List[str] = []

    # Collect paragraph text; keep it simple for academic submission.
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    return "\n".join(parts).strip()


def _extract_pdf_text(pdf_path: str, label: str = "input") -> str:
    """
    Extract text from a PDF file using pdfplumber.

    Args:
        pdf_path: Path to the PDF file.
        label: Label for error messages.

    Returns:
        Extracted text.
    """
    try:
        import pdfplumber  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "PDF input support requires 'pdfplumber'. Install with: pip install -r requirements.txt"
        ) from exc

    parts: List[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                parts.append(text)

    return "\n".join(parts).strip()


def write_text_file(path: str, content: str) -> None:
    """
    Write text content to a file (UTF-8), creating parent directories if needed.

    Args:
        path (str): Output file path.
        content (str): Text content to write.

    Raises:
        OSError: If the file cannot be written.
    """
    directory = os.path.dirname(path)
    if directory:
        os.makedirs(directory, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


def write_json_file(path: str, data: Any) -> None:
    """
    Serialize a Python object as pretty-printed JSON and write it to a file.

    Args:
        path (str): Output file path.
        data (Any): Data to serialize.

    Raises:
        OSError: If the file cannot be written.
    """
    json_str = to_pretty_json(data)
    write_text_file(path, json_str)


def to_pretty_json(data: Any) -> str:
    """
    Convert a Python object to a pretty-printed JSON string.

    Args:
        data (Any): Data to serialize.

    Returns:
        str: JSON string.
    """
    return json.dumps(data, indent=2, ensure_ascii=False)


def safe_json_parse(raw_text: str) -> Optional[Any]:
    """
    Attempt to parse JSON from raw LLM output.

    Handles:
    - Leading/trailing whitespace.
    - Markdown code fences (``` or ```json or ```JSON).
    - Preamble text before the opening brace.

    Args:
        raw_text (str): Raw model output.

    Returns:
        Any: Parsed Python object if successful, or None if parsing fails.
        Callers should treat a None return as a parse failure and apply
        strip_non_json() as a recovery step before retrying.
    """
    text = raw_text.strip()

    # FIX: Strip markdown code fences more reliably using a regex rather than
    # splitting on ``` which can pick the wrong block when the model echoes the
    # schema alongside the actual JSON output.
    fence_match = re.search(r"```(?:json|JSON)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if fence_match:
        text = fence_match.group(1).strip()

    # FIX: Return None on failure instead of raw_text.
    # Returning raw_text caused callers to receive a string when they expected
    # a dict, leading to a cryptic AttributeError on .get() downstream.
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return None


def strip_non_json(raw_text: str) -> str:
    """
    Attempt a crude cleanup to extract a balanced JSON object from noisy text.

    Finds the first '{' and then walks forward to find the matching closing '}'
    using a bracket counter, ensuring the extracted span is balanced. This is
    more reliable than rfind('}') when the model appends trailing commentary
    after the JSON block (which would make rfind pick the last brace of the
    commentary rather than the JSON object's closing brace).

    This is a best-effort fallback — use safe_json_parse first.

    Args:
        raw_text (str): Raw text possibly containing a JSON object.

    Returns:
        str: Candidate JSON substring, or original text if no balanced
             brace pair is found.
    """
    text = raw_text.strip()
    start = text.find("{")
    if start == -1:
        return text

    # FIX: Walk forward with a bracket counter instead of using rfind("}").
    # rfind picks the very last "}" in the string, which may be inside trailing
    # commentary the model added after the JSON block, producing invalid JSON.
    depth = 0
    in_string = False
    escape_next = False

    for i, ch in enumerate(text[start:], start=start):
        if escape_next:
            escape_next = False
            continue
        if ch == "\\" and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start : i + 1]

    # No balanced closing brace found — return from start to end as best effort.
    return text[start:]


def export_cv_docx_and_pdf(
    cv_text: str,
    docx_path: str,
    pdf_path: str,
) -> None:
    """
    Export a plain-text CV into a DOCX file, and (best-effort) convert it to PDF.

    DOCX is always attempted using `python-docx`.
    PDF conversion is attempted only if `pandoc` is available on the system PATH.

    Args:
        cv_text: Plain text CV content.
        docx_path: Output DOCX file path.
        pdf_path: Output PDF file path.

    Raises:
        RuntimeError: If DOCX export fails.
    """
    if not cv_text or not cv_text.strip():
        raise ValueError("CV text is empty; cannot export documents.")

    _export_cv_to_docx(cv_text=cv_text, docx_path=docx_path)
    _export_docx_to_pdf_best_effort(docx_path=docx_path, pdf_path=pdf_path)


def _export_cv_to_docx(cv_text: str, docx_path: str) -> None:
    """
    Write a DOCX document from CV plain text.

    The function treats these as:
    - Section headings: lines exactly equal to one of the standard headings.
    - Bullets: lines starting with '*' or '-'.
    - Everything else: plain paragraphs.
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "DOCX export requires 'python-docx'. Install dependencies with: pip install -r requirements.txt"
        ) from exc

    # Keep the document simple and ATS-friendly.
    doc = Document()

    # Slightly larger readability defaults (not critical for ATS).
    style = doc.styles["Normal"]
    try:
        style.font.size = Pt(11)  # type: ignore[attr-defined]
    except Exception:
        # If font size setting fails, proceed with default.
        pass

    headings = {"SUMMARY", "EDUCATION", "EXPERIENCE", "PROJECTS", "SKILLS", "ACHIEVEMENTS"}

    lines = cv_text.splitlines()
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        upper = stripped.upper()
        if upper in headings:
            doc.add_heading(stripped, level=1)
            continue

        # Bullets
        if stripped.startswith("* ") or stripped.startswith("- "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(bullet_text, style="List Bullet")
            # Ensure no extra trailing spaces.
            _ = p
            continue

        # Default: paragraph
        doc.add_paragraph(stripped)

    # Ensure parent directory exists.
    directory = os.path.dirname(docx_path)
    if directory:
        os.makedirs(directory, exist_ok=True)

    doc.save(docx_path)


def _export_docx_to_pdf_best_effort(docx_path: str, pdf_path: str) -> None:
    """
    Convert DOCX to PDF using Pandoc if available.

    If pandoc is missing, the function logs a warning and skips PDF output.
    """
    # Import lazily so the module works without pandoc installed.
    import subprocess

    # If pandoc isn't present, skip with a warning.
    try:
        subprocess.run(["pandoc", "--version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
    except FileNotFoundError:
        log_warning("Pandoc not found. Skipping PDF export; DOCX output is still generated.")
        return

    try:
        subprocess.run(
            ["pandoc", docx_path, "-o", pdf_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
        )
    except Exception as exc:  # noqa: BLE001
        log_warning(f"PDF export via pandoc failed ({exc}); skipping PDF output.")
        return